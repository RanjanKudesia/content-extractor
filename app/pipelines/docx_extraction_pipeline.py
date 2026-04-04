"""DOCX extraction pipeline that produces structured JSON output."""

import base64
import logging
import mimetypes
from io import BytesIO
from typing import Any
from zipfile import is_zipfile

from docx import Document
from docx.document import Document as DocumentObject
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree

XML_NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}


class DocxExtractionPipeline:
    """Extract paragraphs, tables, styles, numbering, and media from DOCX files."""

    def __init__(self) -> None:
        self.logger = logging.getLogger(__name__)

    def run(self, file_bytes: bytes, output_basename: str) -> tuple[dict[str, Any], str]:
        """Parse a DOCX byte stream and persist the extracted JSON payload."""
        if not is_zipfile(BytesIO(file_bytes)):
            raise ValueError(
                "Invalid DOCX file: not a valid ZIP archive. File may be corrupted.")

        try:
            document = Document(BytesIO(file_bytes))
        except (ValueError, TypeError, OSError, KeyError) as e:
            raise ValueError(
                f"Failed to parse DOCX document: {str(e)}") from e

        extracted = self._extract_document(document, output_basename)
        return extracted, f"virtual://extracted/{output_basename}.json"

    def _extract_document(self, document: DocumentObject, output_basename: str) -> dict[str, Any]:
        media_index = self._extract_and_save_media(document, output_basename)

        paragraphs = [
            self._extract_paragraph(paragraph, index, media_index, document)
            for index, paragraph in enumerate(document.paragraphs)
        ]
        tables = [
            self._extract_table(table, index, media_index, document)
            for index, table in enumerate(document.tables)
        ]

        body_order: list[dict[str, Any]] = []
        paragraph_index = 0
        table_index = 0
        for child in document.element.body.iterchildren():
            tag = child.tag.rsplit("}", 1)[-1]
            if tag == "p":
                body_order.append(
                    {"type": "paragraph", "index": paragraph_index})
                paragraph_index += 1
            elif tag == "tbl":
                body_order.append({"type": "table", "index": table_index})
                table_index += 1

        core = document.core_properties
        sections = []
        for index, section in enumerate(document.sections):
            sections.append(
                {
                    "index": index,
                    "page_width_twips": section.page_width.twips if section.page_width else None,
                    "page_height_twips": section.page_height.twips if section.page_height else None,
                    "left_margin_twips": section.left_margin.twips if section.left_margin else None,
                    "right_margin_twips": section.right_margin.twips if section.right_margin else None,
                    "top_margin_twips": section.top_margin.twips if section.top_margin else None,
                    "bottom_margin_twips": section.bottom_margin.twips if section.bottom_margin else None,
                    "header_distance_twips": section.header_distance.twips if section.header_distance else None,
                    "footer_distance_twips": section.footer_distance.twips if section.footer_distance else None,
                    "start_type": section.start_type.name if section.start_type else None,
                    "different_first_page_header_footer": section.different_first_page_header_footer,
                    "header_paragraphs": [
                        self._extract_paragraph(p, i, media_index, document)
                        for i, p in enumerate(section.header.paragraphs)
                    ],
                    "footer_paragraphs": [
                        self._extract_paragraph(p, i, media_index, document)
                        for i, p in enumerate(section.footer.paragraphs)
                    ],
                }
            )

        return {
            "metadata": {
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "inline_shape_count": len(document.inline_shapes),
                "section_count": len(document.sections),
                "style_count": len(document.styles),
                "core_properties": {
                    "author": core.author,
                    "category": core.category,
                    "comments": core.comments,
                    "content_status": core.content_status,
                    "created": core.created.isoformat() if core.created else None,
                    "identifier": core.identifier,
                    "keywords": core.keywords,
                    "language": core.language,
                    "last_modified_by": core.last_modified_by,
                    "last_printed": core.last_printed.isoformat() if core.last_printed else None,
                    "modified": core.modified.isoformat() if core.modified else None,
                    "revision": core.revision,
                    "subject": core.subject,
                    "title": core.title,
                    "version": core.version,
                },
            },
            "document_defaults": self._extract_document_defaults(document),
            "document_order": body_order,
            "styles": self._extract_styles(document),
            "numbering": self._extract_numbering(document),
            "sections": sections,
            "media": list(media_index.values()),
            "paragraphs": paragraphs,
            "tables": tables,
        }

    def _extract_document_defaults(self, document: DocumentObject) -> dict[str, Any]:
        """Extract DOCX-level default run properties (docDefaults), including theme-resolved font/color."""
        defaults = {
            "font_name": None,
            "font_size_pt": None,
            "color_rgb": None,
        }

        try:
            styles_root = document.styles.element
            run_defaults = styles_root.find(
                "w:docDefaults/w:rPrDefault/w:rPr", XML_NS)
            if run_defaults is None:
                return defaults

            theme_data = self._extract_theme_data(document)

            r_fonts = run_defaults.find("w:rFonts", XML_NS)
            if r_fonts is not None:
                font_name = (
                    r_fonts.get(qn("w:ascii"))
                    or r_fonts.get(qn("w:hAnsi"))
                    or r_fonts.get(qn("w:cs"))
                )
                if font_name is None:
                    theme_font_key = (
                        r_fonts.get(qn("w:asciiTheme"))
                        or r_fonts.get(qn("w:hAnsiTheme"))
                        or r_fonts.get(qn("w:csTheme"))
                        or r_fonts.get(qn("w:eastAsiaTheme"))
                    )
                    font_name = theme_data["fonts"].get(theme_font_key)
                defaults["font_name"] = font_name

            sz_elem = run_defaults.find("w:sz", XML_NS)
            if sz_elem is not None:
                sz_val = sz_elem.get(qn("w:val"))
                if sz_val is not None:
                    defaults["font_size_pt"] = int(sz_val) / 2.0

            color_elem = run_defaults.find("w:color", XML_NS)
            if color_elem is not None:
                color_val = color_elem.get(qn("w:val"))
                if color_val and color_val.lower() != "auto":
                    defaults["color_rgb"] = color_val.upper()
                elif color_elem.get(qn("w:themeColor")):
                    theme_color_key = color_elem.get(qn("w:themeColor"))
                    defaults["color_rgb"] = theme_data["colors"].get(
                        theme_color_key)

        except (AttributeError, KeyError, ValueError, TypeError, etree.XMLSyntaxError) as e:
            self.logger.warning(
                "Failed to extract document defaults", extra={"error": str(e)})

        return defaults

    def _extract_theme_data(self, document: DocumentObject) -> dict[str, dict[str, str]]:
        """Extract major/minor Latin theme fonts and color scheme values from theme part."""
        theme_fonts: dict[str, str] = {}
        theme_colors: dict[str, str] = {}

        try:
            theme_blob = None
            for rel in document.part.rels.values():
                if rel.reltype.endswith("/theme"):
                    theme_blob = rel.target_part.blob
                    break

            if not theme_blob:
                return {"fonts": theme_fonts, "colors": theme_colors}

            root = etree.fromstring(theme_blob)
            ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}

            major_latin = root.find(
                ".//a:themeElements/a:fontScheme/a:majorFont/a:latin", ns)
            minor_latin = root.find(
                ".//a:themeElements/a:fontScheme/a:minorFont/a:latin", ns)

            if major_latin is not None and major_latin.get("typeface"):
                typeface = major_latin.get("typeface")
                theme_fonts["majorAscii"] = typeface
                theme_fonts["majorHAnsi"] = typeface
                theme_fonts["majorBidi"] = typeface

            if minor_latin is not None and minor_latin.get("typeface"):
                typeface = minor_latin.get("typeface")
                theme_fonts["minorAscii"] = typeface
                theme_fonts["minorHAnsi"] = typeface
                theme_fonts["minorBidi"] = typeface

            color_scheme = root.find(".//a:themeElements/a:clrScheme", ns)
            if color_scheme is not None:
                for color_node in color_scheme:
                    tag_name = color_node.tag.rsplit("}", 1)[-1]
                    color_value = None
                    srgb = color_node.find("a:srgbClr", ns)
                    if srgb is not None:
                        color_value = srgb.get("val")
                    else:
                        sys_clr = color_node.find("a:sysClr", ns)
                        if sys_clr is not None:
                            color_value = sys_clr.get("lastClr")

                    if color_value:
                        theme_colors[tag_name] = color_value.upper()

            # Common Word mappings between run themeColor values and clrScheme keys.
            if "dk1" in theme_colors:
                theme_colors["text1"] = theme_colors["dk1"]
            if "dk2" in theme_colors:
                theme_colors["text2"] = theme_colors["dk2"]
            if "lt1" in theme_colors:
                theme_colors["background1"] = theme_colors["lt1"]
            if "lt2" in theme_colors:
                theme_colors["background2"] = theme_colors["lt2"]
            if "hlink" in theme_colors:
                theme_colors["hyperlink"] = theme_colors["hlink"]
            if "folHlink" in theme_colors:
                theme_colors["followedHyperlink"] = theme_colors["folHlink"]

        except (AttributeError, KeyError, ValueError, TypeError, etree.XMLSyntaxError) as e:
            self.logger.warning(
                "Failed to extract theme data", extra={"error": str(e)})

        return {"fonts": theme_fonts, "colors": theme_colors}

    def _extract_and_save_media(self, document: DocumentObject, output_basename: str) -> dict[str, dict[str, Any]]:
        media_index: dict[str, dict[str, Any]] = {}

        for rel_id, rel in document.part.rels.items():
            if rel.reltype != RT.IMAGE:
                continue

            try:
                image_part = rel.target_part
                if image_part is None or not hasattr(image_part, "blob"):
                    continue

                blob = image_part.blob
                if not blob:
                    continue

                content_type = image_part.content_type or "application/octet-stream"
                extension = self._content_type_to_extension(content_type)
                file_name = f"{output_basename}_{rel_id}.{extension}"
                media_index[rel_id] = {
                    "relationship_id": rel_id,
                    "content_type": content_type,
                    "file_name": file_name,
                    "local_file_path": None,
                    "local_url": None,
                    "base64": base64.b64encode(blob).decode("ascii"),
                    "size_bytes": len(blob),
                    "source_partname": str(image_part.partname),
                    "alt_text": None,
                }
            except (AttributeError, KeyError, ValueError, TypeError, OSError) as e:
                self.logger.warning(
                    "Failed to extract media",
                    extra={"rel_id": rel_id, "error": str(e)},
                )
                continue

        return media_index

    def _content_type_to_extension(self, content_type: str) -> str:
        """Convert MIME type to file extension using mimetypes library.
        Handles legacy formats like wmf/emf and falls back to common image formats."""
        normalized = content_type.lower().strip()
        if not normalized:
            return "bin"

        # Try mimetypes library first (supports legacy formats like wmf, emf)
        ext = mimetypes.guess_extension(normalized, strict=False)
        if ext:
            return ext.lstrip(".")

        # Fallback mapping for known formats
        fallback = {
            "image/jpeg": "jpg",
            "image/jpg": "jpg",
            "image/png": "png",
            "image/gif": "gif",
            "image/bmp": "bmp",
            "image/tiff": "tiff",
            "image/webp": "webp",
            "image/svg+xml": "svg",
            "image/x-wmf": "wmf",
            "image/x-emf": "emf",
        }
        return fallback.get(normalized, "bin")

    def _extract_styles(self, document: DocumentObject) -> list[dict[str, Any]]:
        styles_data: list[dict[str, Any]] = []
        for style in document.styles:
            font = getattr(style, "font", None)
            p_format = getattr(style, "paragraph_format", None)
            styles_data.append(
                {
                    "style_id": style.style_id,
                    "name": style.name,
                    "type": str(style.type),
                    "builtin": style.builtin,
                    "hidden": style.hidden,
                    "priority": style.priority,
                    "quick_style": style.quick_style,
                    "font": {
                        "name": font.name if font else None,
                        "size_pt": font.size.pt if font and font.size else None,
                        "bold": bool(font.bold) if font and font.bold is not None else None,
                        "italic": bool(font.italic) if font and font.italic is not None else None,
                        "underline": bool(font.underline) if font and font.underline is not None else None,
                        "all_caps": bool(font.all_caps) if font and font.all_caps is not None else None,
                        "small_caps": bool(font.small_caps) if font and font.small_caps is not None else None,
                        "color_rgb": str(font.color.rgb) if font and font.color and font.color.rgb else None,
                        "highlight_color": font.highlight_color.name if font and font.highlight_color else None,
                    },
                    "paragraph": {
                        "alignment": p_format.alignment.name if p_format and p_format.alignment else None,
                        "left_indent_pt": p_format.left_indent.pt if p_format and p_format.left_indent else None,
                        "right_indent_pt": p_format.right_indent.pt if p_format and p_format.right_indent else None,
                        "first_line_indent_pt": p_format.first_line_indent.pt if p_format and p_format.first_line_indent else None,
                        "space_before_pt": p_format.space_before.pt if p_format and p_format.space_before else None,
                        "space_after_pt": p_format.space_after.pt if p_format and p_format.space_after else None,
                        "line_spacing": p_format.line_spacing if p_format else None,
                        "line_spacing_rule": p_format.line_spacing_rule.name if p_format and p_format.line_spacing_rule else None,
                    },
                }
            )
        return styles_data

    def _extract_numbering(self, document: DocumentObject) -> list[dict[str, Any]]:
        numbering_data: list[dict[str, Any]] = []
        numbering_part = getattr(document.part, "numbering_part", None)
        if numbering_part is None:
            return numbering_data

        root = numbering_part.element
        for num in root.findall("w:num", XML_NS):
            num_id = num.get(qn("w:numId"))
            abstract = num.find("w:abstractNumId", XML_NS)
            numbering_data.append(
                {
                    "num_id": int(num_id) if num_id is not None else None,
                    "abstract_num_id": int(abstract.get(qn("w:val")))
                    if abstract is not None and abstract.get(qn("w:val")) is not None
                    else None,
                }
            )

        return numbering_data

    def _extract_table(
        self,
        table: Table,
        index: int,
        media_index: dict[str, dict[str, Any]],
        document: DocumentObject | None = None,
    ) -> dict[str, Any]:
        rows_data: list[dict[str, Any]] = []
        for row_index, row in enumerate(table.rows):
            cells_data: list[dict[str, Any]] = []
            for cell_index, cell in enumerate(row.cells):
                cells_data.append(self._extract_cell(
                    cell, row_index, cell_index, media_index, document))
            rows_data.append({"row_index": row_index, "cells": cells_data})

        return {
            "index": index,
            "row_count": len(table.rows),
            "column_count": len(table.columns) if table.rows else 0,
            "style": table.style.name if table.style else None,
            "autofit": table.autofit,
            "alignment": table.alignment.name if table.alignment else None,
            "direction": table.table_direction.name if table.table_direction else None,
            "rows": rows_data,
        }

    def _extract_cell(
        self,
        cell: _Cell,
        row_index: int,
        cell_index: int,
        media_index: dict[str, dict[str, Any]],
        document: DocumentObject | None = None,
    ) -> dict[str, Any]:
        """Extract cell content, including nested paragraphs and tables."""
        paragraphs = []
        tables = []
        element_index = 0

        for child in self._element_of(cell).iterchildren():
            local_tag = child.tag.rsplit(
                "}", 1)[-1] if "}" in child.tag else child.tag

            if local_tag == "p":
                para_obj = Paragraph(child, cell)
                para_data = self._extract_paragraph(
                    para_obj, element_index, media_index, document)
                paragraphs.append(para_data)
                element_index += 1

            elif local_tag == "tbl":
                table_obj = Table(child, cell)
                table_data = self._extract_table(
                    table_obj, element_index, media_index)
                tables.append(table_data)
                element_index += 1

        return {
            "row_index": row_index,
            "cell_index": cell_index,
            "text": cell.text,
            "paragraphs": paragraphs,
            "tables": tables,
        }

    def _extract_paragraph(
        self,
        paragraph: Paragraph,
        index: int,
        media_index: dict[str, dict[str, Any]],
        document: DocumentObject | None = None,
    ) -> dict[str, Any]:
        style_name = paragraph.style.name if paragraph.style else None
        paragraph_element = self._element_of(paragraph)
        p_pr = paragraph_element.pPr
        num_pr = p_pr.numPr if p_pr is not None else None

        list_info = None
        if num_pr is not None:
            num_id = num_pr.numId.val if num_pr.numId is not None else None
            ilvl = num_pr.ilvl.val if num_pr.ilvl is not None else None
            list_info = {
                "num_id": int(num_id) if num_id is not None else None,
                "level": int(ilvl) if ilvl is not None else None,
            }

        is_bullet = bool(
            style_name and "bullet" in style_name.lower()) or list_info is not None
        is_numbered = bool(style_name and "number" in style_name.lower())

        tab_stops = []
        for tab in paragraph.paragraph_format.tab_stops:
            tab_stops.append(
                {
                    "position_pt": tab.position.pt if tab.position else None,
                    "alignment": tab.alignment.name if tab.alignment else None,
                    "leader": tab.leader.name if tab.leader else None,
                }
            )

        return_data = {
            "index": index,
            "text": paragraph.text,
            "alignment": paragraph.alignment.name if paragraph.alignment else None,
            "style": style_name,
            "is_heading": bool(style_name and style_name.lower().startswith("heading")),
            "is_bullet": is_bullet,
            "is_numbered": is_numbered,
            "list_info": list_info,
            "numbering_format": None,
            "left_indent_pt": paragraph.paragraph_format.left_indent.pt if paragraph.paragraph_format.left_indent else None,
            "right_indent_pt": paragraph.paragraph_format.right_indent.pt if paragraph.paragraph_format.right_indent else None,
            "first_line_indent_pt": paragraph.paragraph_format.first_line_indent.pt
            if paragraph.paragraph_format.first_line_indent
            else None,
            "line_spacing": paragraph.paragraph_format.line_spacing,
            "line_spacing_rule": paragraph.paragraph_format.line_spacing_rule.name
            if paragraph.paragraph_format.line_spacing_rule
            else None,
            "space_before_pt": paragraph.paragraph_format.space_before.pt if paragraph.paragraph_format.space_before else None,
            "space_after_pt": paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else None,
            "keep_together": paragraph.paragraph_format.keep_together,
            "keep_with_next": paragraph.paragraph_format.keep_with_next,
            "page_break_before": paragraph.paragraph_format.page_break_before,
            "widow_control": paragraph.paragraph_format.widow_control,
            "tab_stops": tab_stops,
            "raw_xml": self._xml_of(paragraph),
            "runs": self._extract_paragraph_runs(paragraph, media_index),
        }

        # Resolve numbering format if document is available and paragraph has list info
        if document is not None and return_data.get("list_info"):
            self._resolve_list_formatting(return_data, document)

        return return_data

    def _extract_paragraph_runs(
        self,
        paragraph: Paragraph,
        media_index: dict[str, dict[str, Any]],
    ) -> list[dict[str, Any]]:
        """Extract all runs from paragraph, including nested children inside hyperlinks."""
        runs_data: list[dict[str, Any]] = []
        run_index = 0

        for child in self._element_of(paragraph).iterchildren():
            local_tag = child.tag.rsplit(
                "}", 1)[-1] if "}" in child.tag else child.tag

            if local_tag == "r":
                run_obj = Run(child, paragraph)
                run_data = self._extract_run(
                    run_obj, run_index, media_index, paragraph)
                runs_data.append(run_data)
                run_index += 1

            elif local_tag == "hyperlink":
                r_id = child.get(qn("r:id"))
                url: str | None = None
                if r_id:
                    try:
                        url = paragraph.part.rels[r_id].target_url
                    except (KeyError, AttributeError):
                        pass

                # Recursively extract all runs inside hyperlink, including nested children.
                for sub_child in child.iterchildren():
                    sub_tag = sub_child.tag.rsplit(
                        "}", 1)[-1] if "}" in sub_child.tag else sub_child.tag
                    if sub_tag == "r":
                        try:
                            run_obj = Run(sub_child, paragraph)
                            run_data = self._extract_run(
                                run_obj, run_index, media_index, paragraph)
                            run_data["hyperlink_url"] = url
                            runs_data.append(run_data)
                            run_index += 1
                        except (AttributeError, ValueError, TypeError) as e:
                            self.logger.warning(
                                "Failed to extract run in hyperlink",
                                extra={"error": str(e)},
                            )
                            continue

        return runs_data

    def _get_drawing_alt_text(self, blip_elem: Any) -> str | None:
        """Walk up from a blip element to find wp:docPr and extract alt text."""
        wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        elem = blip_elem.getparent()
        while elem is not None:
            local = elem.tag.rsplit(
                "}", 1)[-1] if "}" in elem.tag else elem.tag
            if local in ("inline", "anchor"):
                doc_pr = elem.find(f"{{{wp_ns}}}docPr")
                if doc_pr is not None:
                    return doc_pr.get("descr") or doc_pr.get("name") or None
                break
            elem = elem.getparent()
        return None

    def _get_drawing_extents(self, blip_elem: Any) -> tuple[int | None, int | None]:
        """Walk up from a blip element to find the enclosing wp:inline/wp:anchor
        and return its (width_emu, height_emu) from the extent attribute."""
        wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        elem = blip_elem.getparent()
        while elem is not None:
            local = elem.tag.rsplit(
                "}", 1)[-1] if "}" in elem.tag else elem.tag
            if local in ("inline", "anchor"):
                extent = elem.find(f"{{{wp_ns}}}extent")
                if extent is not None:
                    cx = extent.get("cx")
                    cy = extent.get("cy")
                    return (
                        int(cx) if cx is not None else None,
                        int(cy) if cy is not None else None,
                    )
                break
            elem = elem.getparent()
        return None, None

    def _resolve_list_formatting(self, paragraph_data: dict[str, Any], document: DocumentObject) -> None:
        """Resolve abstract numbering format to human-readable list formatting."""
        if not paragraph_data.get("list_info"):
            return

        list_info = paragraph_data["list_info"]
        num_id = list_info.get("num_id")
        ilvl = list_info.get("level")

        if num_id is None:
            return

        try:
            numbering_part = document.part.numbering_part
            if numbering_part is None:
                return

            root = numbering_part.element
            # Find the num element with matching numId
            for num in root.findall("w:num", XML_NS):
                if int(num.get(qn("w:numId")) or -1) == num_id:
                    abstract_elem = num.find("w:abstractNumId", XML_NS)
                    if abstract_elem is not None:
                        abstract_id = int(abstract_elem.get(qn("w:val")) or -1)
                        if abstract_id >= 0:
                            paragraph_data["numbering_format"] = self._get_numbering_format(
                                root, abstract_id, ilvl or 0
                            )
                    break
        except (AttributeError, ValueError, TypeError, KeyError) as e:
            self.logger.warning(
                "Failed to resolve list formatting", extra={"error": str(e)})

    def _get_numbering_format(self, numbering_root: Any, abstract_id: int, level: int) -> str | None:
        """Extract the numbering format string from the abstract numbering definition."""
        for abs_num in numbering_root.findall("w:abstractNum", XML_NS):
            if int(abs_num.get(qn("w:abstractNumId")) or -1) == abstract_id:
                # Find the level definition for this level
                for lvl in abs_num.findall("w:lvl", XML_NS):
                    if int(lvl.get(qn("w:ilvl")) or -1) == level:
                        # Extract numFmt and lvlText
                        num_fmt = lvl.find("w:numFmt", XML_NS)
                        lvl_text = lvl.find("w:lvlText", XML_NS)
                        if num_fmt is not None and lvl_text is not None:
                            fmt = num_fmt.get(qn("w:val"))
                            text = lvl_text.get(qn("w:val"))
                            return f"{fmt}:{text}"
        return None

    def _resolve_run_font_properties(
        self,
        run: Run,
        paragraph: Paragraph | None = None,
    ) -> tuple[str | None, float | None]:
        """
        Resolve font name and size for a run, considering:
        1. Explicit run-level formatting
        2. Paragraph style formatting
        3. Parent style formatting (recursively)

        Returns: (font_name, font_size_pt)
        """
        # First, check explicit run-level formatting
        if run.font.name is not None:
            run_font_name = run.font.name
        else:
            run_font_name = None

        if run.font.size is not None:
            run_font_size_pt = run.font.size.pt
        else:
            run_font_size_pt = None

        # If we have explicit run-level formatting, use it
        if run_font_name is not None and run_font_size_pt is not None:
            return run_font_name, run_font_size_pt

        # Otherwise, walk up the style hierarchy to find the values
        if paragraph is None:
            # Can't resolve from style without paragraph reference
            return run_font_name, run_font_size_pt

        try:
            style = paragraph.style
            if style is None:
                return run_font_name, run_font_size_pt

            # Walk the style hierarchy
            visited_styles = set()
            current_style = style
            style_font_name = None
            style_font_size_pt = None

            while current_style is not None and current_style.style_id not in visited_styles:
                visited_styles.add(current_style.style_id)

                # Check this style's font properties
                if style_font_name is None and current_style.font.name is not None:
                    style_font_name = current_style.font.name

                if style_font_size_pt is None and current_style.font.size is not None:
                    style_font_size_pt = current_style.font.size.pt

                # Move to parent style
                current_style = current_style.base_style

            # Use resolved style values if run didn't have them
            final_font_name = run_font_name or style_font_name
            final_font_size_pt = run_font_size_pt or style_font_size_pt

            return final_font_name, final_font_size_pt

        except (AttributeError, ValueError, TypeError, KeyError) as e:
            self.logger.warning(
                "Failed to resolve font properties from style",
                extra={"error": str(e)},
            )
            return run_font_name, run_font_size_pt

    def _extract_run(
        self,
        run: Run,
        index: int,
        media_index: dict[str, dict[str, Any]],
        paragraph: Paragraph | None = None,
    ) -> dict[str, Any]:
        font = run.font
        color = font.color.rgb if font.color is not None else None

        embedded_media = []
        drawing_blips = self._element_of(
            run).xpath(".//*[local-name()='blip']")
        for blip in drawing_blips:
            rel_id = blip.get(qn("r:embed"))
            if rel_id and rel_id in media_index:
                entry = dict(media_index[rel_id])
                width_emu, height_emu = self._get_drawing_extents(blip)
                entry["width_emu"] = width_emu
                entry["height_emu"] = height_emu
                embedded_media.append(entry)

        # Resolve font name and size from run, style hierarchy, or defaults
        resolved_font_name, resolved_font_size = self._resolve_run_font_properties(
            run, paragraph
        )

        return {
            "index": index,
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "strike": bool(run.font.strike) if run.font.strike is not None else False,
            "double_strike": bool(run.font.double_strike) if run.font.double_strike is not None else False,
            "subscript": bool(run.font.subscript) if run.font.subscript is not None else False,
            "superscript": bool(run.font.superscript) if run.font.superscript is not None else False,
            "rtl": bool(run.font.rtl) if run.font.rtl is not None else False,
            "all_caps": bool(font.all_caps) if font.all_caps is not None else False,
            "small_caps": bool(font.small_caps) if font.small_caps is not None else False,
            "font_name": resolved_font_name,
            "font_size_pt": resolved_font_size,
            "color_rgb": str(color) if color is not None else None,
            "highlight_color": font.highlight_color.name if font.highlight_color else None,
            "raw_xml": self._xml_of(run),
            "embedded_media": embedded_media,
        }

    @staticmethod
    def _element_of(node: Any) -> Any:
        return getattr(node, "_element")

    @staticmethod
    def _xml_of(node: Any) -> str:
        return DocxExtractionPipeline._element_of(node).xml
